"""
Extract DOCX content to Markdown + images.

Usage:
    python utils/extract_docx.py <input.docx> [output_dir]

Output structure:
    output_dir/
        index.md          — full markdown with document content
        images/
            img_0.png
            ...
"""

import argparse
import re
import sys
from pathlib import Path

# Fix Windows console encoding for Cyrillic paths
if sys.stdout.encoding and sys.stdout.encoding.lower() not in ("utf-8", "utf-8-sig"):
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile


# ── Heading detection ─────────────────────────────────────────────────────────

HEADING_STYLES = {
    "heading 1": "#",
    "heading 2": "##",
    "heading 3": "###",
    "heading 4": "####",
    "heading 5": "#####",
    "title": "#",
    "subtitle": "##",
    # Russian style names
    "заголовок 1": "#",
    "заголовок 2": "##",
    "заголовок 3": "###",
    "заголовок 4": "####",
}


def para_to_md(para) -> str:
    """Convert a single paragraph to a markdown line."""
    text = para.text.strip()
    if not text:
        return ""

    style_name = para.style.name.lower() if para.style and para.style.name else ""

    # Heading
    prefix = HEADING_STYLES.get(style_name)
    if prefix:
        return f"{prefix} {text}"

    # List paragraph
    num_xml = para._p.find(qn("w:numPr"))
    if num_xml is not None:
        level_el = num_xml.find(qn("w:ilvl"))
        level = int(level_el.get(qn("w:val"), 0)) if level_el is not None else 0
        indent = "  " * level
        return f"{indent}- {text}"

    # Bold-only paragraph → treat as bold text
    if all(run.bold for run in para.runs if run.text.strip()):
        return f"**{text}**"

    return text


def extract_images_from_docx(docx_path: Path, images_dir: Path) -> dict[str, str]:
    """
    Extract embedded images from docx ZIP and return a mapping
    { relationship_id -> relative_path_in_md }.
    """
    images_dir.mkdir(parents=True, exist_ok=True)
    rel_map: dict[str, str] = {}  # rId -> "images/filename"

    with zipfile.ZipFile(str(docx_path)) as zf:
        names = zf.namelist()

        # Parse word/_rels/document.xml.rels to build rId -> target map
        rels_path = "word/_rels/document.xml.rels"
        if rels_path not in names:
            return rel_map

        import xml.etree.ElementTree as ET
        rels_xml = zf.read(rels_path)
        root = ET.fromstring(rels_xml)
        ns = {"r": "http://schemas.openxmlformats.org/package/2006/relationships"}

        img_counter = 0
        for rel in root.findall("r:Relationship", ns):
            rtype = rel.get("Type", "")
            if "image" not in rtype.lower():
                continue
            target = rel.get("Target", "")
            rid = rel.get("Id", "")
            # target like "media/image1.png"
            zip_path = f"word/{target}" if not target.startswith("/") else target.lstrip("/")
            if zip_path not in names:
                continue

            suffix = Path(target).suffix or ".png"
            img_name = f"img_{img_counter}{suffix}"
            img_path = images_dir / img_name
            img_path.write_bytes(zf.read(zip_path))
            rel_map[rid] = f"images/{img_name}"
            img_counter += 1

    return rel_map


def get_para_image_rids(para) -> list[str]:
    """Return list of relationship IDs for inline images in a paragraph."""
    rids: list[str] = []
    # drawings embed blipFill with r:embed attribute
    ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    for drawing in para._p.iter("{http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing}blipFill"):
        pass  # not needed here

    # Inline images: w:drawing > wp:inline > a:graphic > ... > a:blip r:embed
    a_blip_tag = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    for blip in para._p.iter(a_blip_tag):
        embed = blip.get(f"{{{ns_r}}}embed")
        if embed:
            rids.append(embed)
    return rids


def extract_docx(docx_path: Path, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    images_dir = output_dir / "images"

    doc = Document(str(docx_path))

    # Extract images first to build rId map
    rel_map = extract_images_from_docx(docx_path, images_dir)

    md_lines: list[str] = [f"# {docx_path.stem}\n"]

    for para in doc.paragraphs:
        # Check for inline images in this paragraph
        img_rids = get_para_image_rids(para)
        for rid in img_rids:
            if rid in rel_map:
                img_rel = rel_map[rid]
                img_name = Path(img_rel).name
                md_lines.append(f"![{img_name}]({img_rel})")

        line = para_to_md(para)
        if line:
            md_lines.append(line)
        elif not img_rids:
            # preserve blank lines between blocks (max one)
            if md_lines and md_lines[-1] != "":
                md_lines.append("")

    # Tables
    for table in doc.tables:
        md_lines.append("")
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            md_lines.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                md_lines.append("|" + "|".join(["---"] * len(cells)) + "|")
        md_lines.append("")

    # Remove trailing blank lines
    while md_lines and md_lines[-1] == "":
        md_lines.pop()

    md_path = output_dir / "index.md"
    md_path.write_text("\n".join(md_lines), encoding="utf-8")
    print(f"[docx] Done: {md_path}  (images: {images_dir})")


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract DOCX → Markdown + images")
    parser.add_argument("docx", help="Path to .docx file")
    parser.add_argument("output", nargs="?", help="Output directory (default: <docx_stem>/)")
    args = parser.parse_args()

    docx_path = Path(args.docx)
    if not docx_path.exists():
        print(f"Error: file not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    output_dir = Path(args.output) if args.output else docx_path.parent / docx_path.stem
    extract_docx(docx_path, output_dir)


if __name__ == "__main__":
    main()
